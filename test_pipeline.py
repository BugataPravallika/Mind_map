import os
import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '.'))
sys.path.append(os.path.join(os.path.dirname(__file__), 'ai_study_mapper'))
from ai_study_mapper.src.pipeline import StudyMapPipeline

def test():
    # Create dummy docx file if not exists
    import docx
    test_file = "test_content.docx"
    doc = docx.Document()
    doc.add_heading("The Science of Photosynthesis", 0)
    doc.add_paragraph(
        "Photosynthesis is a complex process used by plants to convert light energy into chemical energy. "
        "This occurs specifically in the chloroplasts, where chlorophyll absorbs sunlight. "
        "The light-dependent reactions happen in the thylakoid membranes, producing ATP and NADPH. "
        "Meanwhile, the Calvin cycle takes place in the stroma to fix carbon dioxide into glucose."
    )
    doc.add_paragraph(
        "Importance: This process is the foundation of the food chain and releases oxygen into the atmosphere. "
        "Without it, life on Earth as we know it would not exist. Students often find the Calvin cycle difficult "
        "to memorize because of the complex enzymatic cycles involve."
    )
    doc.save(test_file)
    
    pipeline = StudyMapPipeline(output_dir="data/test_output")
    try:
        results = pipeline.run(
            file_paths=[test_file],
            target_lang="en",
            available_time_mins=60,
            generate_audio=False
        )
        if results:
            print("Pipeline Success!")
            print("Map Path:", results.get("map_path"))
            print("Quiz Count:", len(results.get("quiz", [])))
            print("Summary Length:", len(results.get("simplified_text", "")))
    except Exception as e:
        print(f"Pipeline Failed: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if os.path.exists(test_file):
            os.remove(test_file)

if __name__ == "__main__":
    test()
